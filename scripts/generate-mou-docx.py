#!/usr/bin/env python3
"""Generate MoU Word document for college partnership."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "MoU_Hindusthan_College_Foundation_Track.docx"
FAVICON = ROOT / "docs" / "assets" / "mantra-ai-brand" / "mantra-ai-favicon.png"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    if bold:
        run.bold = True


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

    if FAVICON.exists():
        logo = doc.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo.add_run().add_picture(str(FAVICON), width=Inches(0.85))
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("MEMORANDUM OF UNDERSTANDING (MoU)")
    t_run.bold = True
    t_run.font.size = Pt(16)
    t_run.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub.add_run(
        "Industry-Partnered Internship Program — Mantra.ai\n"
        "Mantra AI Talent Accelerator Internship Program"
    )
    s_run.font.size = Pt(11)
    s_run.font.name = "Calibri"

    doc.add_paragraph()

    add_para(
        doc,
        "This Memorandum of Understanding (“MoU”) is entered into on this ______ day of "
        "______________ 2026 (“Effective Date”)",
    )
    doc.add_paragraph()

    add_para(doc, "BETWEEN", bold=True)
    add_para(
        doc,
        "Mantra.ai, represented by Mr. Praveen Manoharan, CTO & Technical Program Manager "
        "(hereinafter “Mantra.ai” or the “Service Provider”)",
    )
    doc.add_paragraph()
    add_para(doc, "AND", bold=True)
    add_para(
        doc,
        "Hindusthan College of Engineering and Technology (HiCET), an autonomous institution "
        "affiliated to Anna University, Chennai, approved by AICTE, having its campus at "
        "Valley Campus, Pollachi Highway, Othakkalmandapam, Coimbatore – 641032, Tamil Nadu, "
        "India, represented by Dr. J. Jaya, Principal (hereinafter the “College”)",
    )
    doc.add_paragraph()
    add_para(
        doc,
        "Mantra.ai and the College are individually a “Party” and collectively the “Parties”.",
    )

    add_heading(doc, "1. PURPOSE AND BACKGROUND", 2)
    add_bullet(doc, "Mantra.ai operates an internship learning platform for engineering colleges.")
    add_bullet(doc, "The College seeks industry-aligned internship learning for student employability.")
    add_bullet(
        doc,
        "The Parties collaborate to deliver Foundation Track: Python, Data & AI to eligible College students.",
    )

    add_heading(doc, "2. KEY DEFINITIONS", 2)
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("Program", "Mantra AI Talent Accelerator via Mantra.ai"),
        ("Track", "Foundation Track: Python, Data & AI"),
        ("Course Fee", "₹3,999 per Student per enrollment"),
        ("Student", "Bonafide 2nd–4th year student (CSE, IT, ECE, AI/ML, MCA or as agreed)"),
        ("Platform", "Mantra.ai web app and admin portal"),
    ]
    for i, (k, v) in enumerate(rows_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    add_heading(doc, "3. SCOPE", 2)
    add_bullet(doc, "Mantra.ai provides Platform access and Track delivery.")
    add_bullet(doc, "College promotes Program, facilitates enrollment, and nominates TPO + faculty liaison.")
    add_bullet(
        doc,
        "Joint branding: “Hindusthan College of Engineering and Technology × Mantra.ai Industry Internship Lab”",
    )

    add_heading(doc, "4. PROGRAM — FOUNDATION TRACK: PYTHON, DATA & AI", 2)
    add_para(doc, "Duration: 10 intensive weeks (~6 months structured depth) · Level: Beginner Foundation")
    add_para(
        doc,
        "Coverage: Linux, Git, core Python, DSA, OOP, REST APIs, database engineering, "
        "Python for data science, AI foundations, and capstone sprint.",
    )
    doc.add_paragraph()
    add_para(doc, "Per-student deliverables include:", bold=True)
    deliverables = [
        "Lifetime platform access for enrolled Track",
        "200+ lessons, browser-based coding labs (Monaco editor)",
        "Weekly live mentor sessions with recordings",
        "24/7 lesson-scoped AI tutor (disabled during quizzes)",
        "50-question weekly quizzes (80% pass, up to 3 retries)",
        "Interview prep modules per week",
        "Mini projects + capstone + GitHub portfolio",
        "XP, leaderboard, QR-verified certificate",
        "Mantra.ai Hackathon access",
        "Merit-based LoR eligibility",
        "Top 5%: internship pathway (performance-based)",
    ]
    for d in deliverables:
        add_bullet(doc, d)

    add_heading(doc, "5. ROLES AND RESPONSIBILITIES", 2)
    add_para(doc, "Mantra.ai shall:", bold=True)
    for item in [
        "Host Platform; provide curriculum updates",
        "Weekly live sessions + recordings",
        "Admin training, bulk upload template, kickoff webinar",
        "Analytics, certificates, hackathon facilitation",
        "Revenue share settlement per Clause 6",
    ]:
        add_bullet(doc, item)

    add_para(doc, "College shall:", bold=True)
    for item in [
        "Nominate TPO champion and faculty liaison",
        "Ensure the College Coordinator (TPO/Internship Coordinator) provides one official institutional email ID of the College (e.g., @hicet.ac.in) for formal Program communication",
        "Promote to eligible students; facilitate enrollment/fees",
        "Ensure laptop, browser, and internet access",
        "Not redistribute Platform content",
        "Communicate student grievances within reasonable time",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5.3 OFFICIAL COMMUNICATION CHANNEL", 2)
    add_para(
        doc,
        "Mantra.ai shall designate one official company email ID as the sole primary communication "
        "channel for all Program-related operational correspondence with the College (enrollment "
        "support, analytics, mentor sessions, grievance handling, and admin coordination).",
    )
    add_para(doc, "Mantra.ai official communication email: _________________________________ @mantra.ai", bold=True)
    add_para(
        doc,
        "The College Coordinator (TPO/Internship Coordinator) shall use this Mantra.ai official email "
        "for all routine Program communication and shall provide their official College email ID "
        "(as stated above) for formal notices under this MoU.",
    )
    add_para(
        doc,
        "Communication sent to or from the designated official email IDs shall constitute valid "
        "operational correspondence under this MoU, in addition to the contact persons in Clause 16.",
    )

    add_heading(doc, "6. FINANCIAL TERMS", 2)
    add_para(doc, "Course Fee: ₹3,999 per Student (all-inclusive as per Section 4).", bold=True)
    add_para(doc, "College Revenue Share: 30% of Course Fee per confirmed enrollment.", bold=True)

    fin_table = doc.add_table(rows=3, cols=2)
    fin_table.style = "Table Grid"
    fin_table.rows[0].cells[0].text = "Component"
    fin_table.rows[0].cells[1].text = "Amount (per enrollment)"
    fin_table.rows[1].cells[0].text = "College share (30%)"
    fin_table.rows[1].cells[1].text = "₹1,200"
    fin_table.rows[2].cells[0].text = "Mantra.ai share (70%)"
    fin_table.rows[2].cells[1].text = "₹2,799"

    doc.add_paragraph()
    add_para(
        doc,
        "Payment: Students pay Mantra.ai directly (Option A) OR College collects and remits (Option B). "
        "Settlement within 15 business days after month-end. GST as per law. "
        "Refund within 7 days pre-Week 1 on case-by-case basis.",
    )

    add_heading(doc, "7. IMPLEMENTATION TIMELINE", 2)
    timeline = doc.add_table(rows=5, cols=3)
    timeline.style = "Table Grid"
    timeline.rows[0].cells[0].text = "Phase"
    timeline.rows[0].cells[1].text = "Duration"
    timeline.rows[0].cells[2].text = "Activity"
    phases = [
        ("MoU & Setup", "Week 1", "Sign MoU, admin config, bulk upload"),
        ("Orientation", "Week 1–2", "Admin training, kickoff webinar"),
        ("Delivery", "Weeks 1–10", "Modules, live sessions, quizzes, capstone"),
        ("Completion", "Week 10", "Certificates, hackathon, analytics review"),
    ]
    for i, row in enumerate(phases, start=1):
        for j, val in enumerate(row):
            timeline.rows[i].cells[j].text = val

    add_heading(doc, "8–15. IP, DATA, CONFIDENTIALITY, TERM, LIABILITY, DISPUTES", 2)
    add_bullet(doc, "Curriculum and Platform IP remain with Mantra.ai; students own their code/projects.")
    add_bullet(doc, "Student data used only for delivery, analytics, certification, placement support.")
    add_bullet(doc, "Term: 12 months, auto-renewable; 30-day notice for non-renewal.")
    add_bullet(doc, "Confidentiality survives 3 years.")
    add_bullet(doc, "Arbitration in Coimbatore under Arbitration and Conciliation Act, 1996.")
    add_bullet(doc, "Placement/internship not guaranteed; merit-based pathways only.")

    add_heading(doc, "16. CONTACT PERSONS", 2)
    contacts = doc.add_table(rows=5, cols=4)
    contacts.style = "Table Grid"
    contacts.rows[0].cells[0].text = "Party"
    contacts.rows[0].cells[1].text = "Name / Role"
    contacts.rows[0].cells[2].text = "Email"
    contacts.rows[0].cells[3].text = "Phone"
    contact_rows = [
        ("Mantra.ai", "Praveen Manoharan, CTO & TPM", "praveen.manoharan@mantra.ai", "+91 8760 8760 62"),
        ("Mantra.ai", "Official communication (Company)", "_________________ @mantra.ai", "—"),
        ("College", "Dr. J. Jaya, Principal", "principal@hicet.ac.in", "0422-4242424"),
        ("College", "TPO / Internship Coordinator", "_________________ @hicet.ac.in", "_________________"),
    ]
    for i, row in enumerate(contact_rows, start=1):
        for j, val in enumerate(row):
            contacts.rows[i].cells[j].text = val

    add_heading(doc, "17. SIGNATURES", 2)
    doc.add_paragraph()
    add_para(doc, "FOR AND ON BEHALF OF MANTRA.AI")
    add_para(doc, "Name: Mr. Praveen Manoharan")
    add_para(doc, "Designation: CTO & Technical Program Manager")
    add_para(doc, "Signature: _________________________    Date: _________________________")
    doc.add_paragraph()
    add_para(doc, "FOR HINDUSTHAN COLLEGE OF ENGINEERING AND TECHNOLOGY (HiCET)")
    add_para(doc, "Name: Dr. J. Jaya")
    add_para(doc, "Designation: Principal")
    add_para(doc, "Signature: _________________________    Date: _________________________")

    add_heading(doc, "ANNEXURE A — IMPLEMENTATION SCHEDULE", 2)
    annex = doc.add_table(rows=10, cols=2)
    annex.style = "Table Grid"
    annex_data = [
        ("Track", "Foundation Track: Python, Data & AI"),
        ("Course Fee", "₹3,999 per Student"),
        ("Cohort size", "__________ Students"),
        ("Departments", "__________"),
        ("Start date", "__________"),
        ("Payment option", "Option A / Option B"),
        ("TPO champion", "Name: __________ Email: __________ @hicet.ac.in"),
        ("Faculty liaison", "Name: __________ Email: __________"),
        ("Mantra.ai official communication email", "__________ @mantra.ai (Clause 5.3)"),
    ]
    for i, (k, v) in enumerate(annex_data):
        annex.rows[i].cells[0].text = k
        annex.rows[i].cells[1].text = v

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer.add_run(
        "Proprietary — Mantra.ai · Draft MoU · June 2026"
    )
    f_run.font.size = Pt(9)
    f_run.font.name = "Calibri"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    build()
